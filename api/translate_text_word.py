import io
import os
import tempfile

import translators as ts
import fitz
from api.s3_client import upload_to_s3
from api.s3_client import s3_client, AWS_BUCKET_NAME
from docx import Document


def split_text(text, max_length=1000):
    """Разделяет текст на части, не превышающие max_length."""
    words = text.split()
    current_part = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 > max_length:
            yield ' '.join(current_part)
            current_part = [word]
            current_length = len(word)
        else:
            current_part.append(word)
            current_length += len(word) + 1

    if current_part:
        yield ' '.join(current_part)


async def detect_and_transl_text_word(file_name: str):
    file_obj = s3_client.get_object(Bucket=AWS_BUCKET_NAME, Key=file_name)
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_file.write(file_obj['Body'].read())
        temp_file_path = temp_file.name

    document = Document(temp_file_path)
    detected_text = '\n'.join([para.text for para in document.paragraphs if para.text.strip()])
    os.remove(temp_file_path)

    translated_parts = []
    for part in split_text(detected_text):
        trans_text = ts.translate_text(query_text=part,
                                       translator='bing',
                                       from_language='auto',
                                       to_language='en')
        translated_parts.append(trans_text)

    final_translated_text = '\n'.join(translated_parts)
    return final_translated_text


async def create_word(query_text: str, file_name: str):
    translated_word_filename = f"trans_{str(file_name)[:-5]}.docx"
    word_buffer = io.BytesIO()

    word_writer = Document()
    word_writer.add_paragraph(query_text)

    word_writer.save(word_buffer)

    word_buffer.seek(0)
    await upload_to_s3(word_buffer, translated_word_filename)
    return translated_word_filename
